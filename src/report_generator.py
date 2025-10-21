import subprocess
import os
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pyquotegen
import streamlit

import ingest

class TreasurerReport:
    GOLD = RGBColor(160, 116, 0)   # Old Gold
    BLACK = RGBColor(34, 31, 32)   # Black

    def __init__(self, start_date, end_date, data_manager: ingest.DataManager):
        """
        data dictionary should include:
        {
            "transaction_notes_table": DataFrame,
            "deposits": DataFrame,
            "deposits_total": float,
            "incomes": DataFrame,
            "withdrawals": DataFrame,
            "withdrawals_total": float,
            "expenses": DataFrame,
            "balance_table": DataFrame,
            "initial_balance": float,
            "current_balance": float,
        }
        """
        self.start_date = start_date
        self.end_date = end_date
        self.data_manager = data_manager
        self.data = data_manager.load_data(start_date,end_date, source="")
        self.dues = data_manager.load_dues(start_date,end_date)
        self.document = Document()

     # ---------------------------
    # Build Report
    # ---------------------------
    def build(self):
        self._set_margins(0.5, 0.5, 0.5, 0.5)  # half-inch margins all around

        # Cover/Header
        self._add_header("./gblogo.png")
        # self._add_heading("ALPHA PHI ALPHA FRATERNITY, INC.", level=0, color=self.GOLD)
        # self._add_heading("Gamma Beta Chapter", level=1, color=self.BLACK)
        # self._add_heading("Treasurer's Report", level=1, color=self.GOLD)
        self._add_heading(
            f"Financial Summary Report from: {self.start_date.strftime('%B %d, %Y')} – {self.end_date.strftime('%B %d, %Y')}",
            level=1,
            color=self.BLACK
        )

        # Financial summary
        # self.document.add_heading("Financial Summary", level=2)
        checking_data = self.data_manager.load_data(self.start_date,self.end_date, source="Checking")
        self.document.add_paragraph(f"Current Balance Checking: ${checking_data['current_balance']:,.2f}")
        self.document.add_paragraph(f"Initial Balance Checking: ${checking_data['initial_balance']:,.2f}")
        cashapp_data = self.data_manager.load_data(self.start_date,self.end_date, source="Cashapp")
        self.document.add_paragraph(f"Current Balance Cashapp: ${cashapp_data['current_balance']:,.2f}")
        self.document.add_paragraph(f"Initial Balance Cashapp: ${cashapp_data['initial_balance']:,.2f}")

        self.document.add_page_break()

        # Receipts (Deposits + Incomes)
        deposits_total  = f"Total Deposits: ${self.data.get('deposits_total'):,.2f}"
        self._add_heading(deposits_total,2,self.GOLD,"left")
        self._add_table(self.data.get("incomes")[["Labels", "Amount"]] , "Income Categories")
        self._add_table(self.data.get("deposits").drop(columns="TransactionId"), "Receipts", )
        self.document.add_page_break()

        # Expenditures (Withdrawals + Expenses)
        withdrawals_total  = f"Total Withdrawals ${self.data.get('withdrawals_total'):,.2f}"

        self._add_heading( withdrawals_total ,2,self.GOLD,"left")
        self._add_table(self.data.get("expenses")[["Labels", "Amount"]], "Expense Categories")
        self._add_table(self.data.get("withdrawals").drop(columns="TransactionId"), "Expenditures")
        self.document.add_page_break()

        # Outstanding Expenditures (placeholder section)
        # self.document.add_heading("Outstanding Expenditures", level=2)
        # self.document.add_paragraph("No outstanding expenditures recorded.")
        # self.document.add_page_break()


        #Dues
        self._add_heading("Dues", level=2, color=self.BLACK, align="left")

        expected_dues = self.dues.get("expected_dues")
        total_paid = self.dues.get("total_paid")
        missed_revenue = self.dues.get("missed_revenue")
        dues_status = self.dues.get("dues_status")
        status_counts = self.dues.get("status_counts")
        status_counts["Count"] = status_counts["Count"].apply(str)

        # Summary text
        if expected_dues is not None:
            self.document.add_paragraph(f"Expected Dues: ${expected_dues:,.2f}")
            self.document.add_paragraph(f"Total Paid: ${total_paid:,.2f}")
            self.document.add_paragraph(f"Missed Revenue: ${missed_revenue:,.2f}")
        else:
            self.document.add_paragraph("No dues data recorded.")

        # Status breakdown table
        if status_counts is not None and not status_counts.empty:
            self._add_table(status_counts, "Member Dues Status Counts")

        # Full per-member dues table
        if dues_status is not None and not dues_status.empty:
            self._add_table(
                dues_status[["MemberName", "PeriodStart", "Amount", "AmountPaid", "Status"]],
                "Per-Member Dues Status"
            )
        
        self._add_footer()

    def save(self):  
        # Save file
        self.filename = (
            f"./reports/automated/TreasurerReport_"
            f"{self.start_date.strftime('%Y-%m-%d')}_to_{self.end_date.strftime('%Y-%m-%d')}.docx"
        )        
        # if os.path.exists(self.filename):
        #     os.remove(self.filename)
        buffer = BytesIO()
        self.document.save(buffer)
        return self.filename, buffer


    def open(self):
        # Auto-open
        subprocess.run(["open", self.filename])

    def _set_margins(self, top=0.5, bottom=0.5, left=0.5, right=0.5):
        """
        Adjusts document margins in inches (default = 0.5").
        Applies to all sections of the document.
        """
        for section in self.document.sections:
            section.top_margin = Inches(top)
            section.bottom_margin = Inches(bottom)
            section.left_margin = Inches(left)
            section.right_margin = Inches(right)
            section.header_distance = Inches(0.1)

    def _add_header(self, logo_path):
        section = self.document.sections[0]
        header = section.header

        # Clear existing header content
        header.is_linked_to_previous = False
        header.paragraphs[0]._element.clear_content()


        # Add logo
        paragraph = header.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(logo_path, width=Inches(.75),)  # adjust size
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add header text below the logo
        paragraph = header.add_paragraph(
            "ALPHA PHI ALPHA FRATERNITY, INC.\n"
            "Gamma Beta Chapter\n"
            "Treasurer's Report\n"
            "Bro. Emmanuel Davis, GB Chapter Treasurer"
        )
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.runs[0]
        run.font.size = Pt(12)
        run.bold = True
        
    def _add_footer(self):
        # Get a random quote from a specific category (e.g., "inspirational")
        inspirational_quote = pyquotegen.get_quote("inspirational")

        footer_text = inspirational_quote + "\n This report has been automatically generated by the Treasury Dashboard ©2025 Terrell Parker "
        
        for section in self.document.sections:
            footer = section.footer
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            run = p.add_run(footer_text)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(150, 150, 150) 
            p.alignment = 1  # center the footer

    # ---------------------------
    # Helpers
    # ---------------------------
    def _add_heading(self, text, level=0, color=None, align="center"):
        para = self.document.add_heading(text, level=level)
        run = para.runs[0]
        run.font.size = Pt(16 if level == 0 else 14)
        if color:
            run.font.color.rgb = color
        if align == "center":
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        return para

    def _add_table(self, df, title, total=None):
        self.document.add_heading(title, level=2)

        # if total is not None:
        #     para = self.document.add_heading(f"Total {title}: ${total:,.2f}", level=3)
        #     para.runs[0].bold = True
        #     para.runs[0].font.color.rgb = self.GOLD


        if df is None or df.empty:
            self.document.add_paragraph(f"No {title.lower()} recorded.")
            return

        table = self.document.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
        table.style = "Light Grid Accent 1"

        # Header row
        for j, col in enumerate(df.columns):
            cell = table.cell(0, j)
            cell.text = col
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.color.rgb = self.GOLD

        # Data rows
        for i in range(df.shape[0]):
            for j in range(df.shape[1]):
                cell = table.cell(i + 1, j)
                value = df.values[i, j]
                if isinstance(value, (float,int)):
                        cell.text = f"${value:,.2f}"
                else:
                    cell.text = str(value)

      
   
