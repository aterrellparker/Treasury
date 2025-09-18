from docx import Document
import subprocess

class TreasurerReport:
    def __init__(self, start_date, end_date, newtable, initial_balance, current_balance, deposits, deposits_sum, withdrawals, withdrawal_sum):
        self.start_date = start_date
        self.end_date = end_date
        self.newtable = newtable
        self.initial_balance = initial_balance
        self.current_balance = current_balance
        self.deposits = deposits
        self.deposits_sum = deposits_sum
        self.withdrawals = withdrawals
        self.withdrawal_sum = withdrawal_sum
        self.document_name = f"TreasurerReport_{end_date.strftime('%B%Y')}.docx"
        self.document = Document()

    def _add_heading(self):
        self.document.add_heading('Alpha Phi Alpha Fraternity, Inc.', level=0)
        self.document.add_heading("Gamma Beta Chapter Treasurer's Report", level=1)
        self.document.add_heading(
            f"{self.start_date.strftime('%B %d, %Y')} - {self.end_date.strftime('%B %d, %Y')}", 
            level=1
        )

    def _add_balances(self):
        self.document.add_paragraph(f"Initial Balance: ${self.initial_balance}")
        self.document.add_paragraph(f"Current Balance: ${self.current_balance}")

    def _add_table(self, title, df, total_label, total_value):
        self.document.add_heading(title, level=2)
        self.document.add_paragraph(f"{total_label}: ${total_value}")

        table = self.document.add_table(df.shape[0] + 1, df.shape[1])
        
        # Header row
        for j, col in enumerate(df.columns):
            table.cell(0, j).text = str(col)

        # Data rows
        for i in range(df.shape[0]):
            for j in range(df.shape[1]):
                table.cell(i + 1, j).text = str(df.iat[i, j])

    def build(self):
        self._add_heading()
        self._add_balances()
        self._add_table("Deposits", self.deposits, "Total Deposits", self.deposits_sum)
        self._add_table("Withdrawals", self.withdrawals, "Total Withdrawals", self.withdrawal_sum)

        self.document.save(self.document_name)
        return self.document_name

    def open(self):
        subprocess.run(['open', self.document_name], capture_output=False)


# === Example Usage ===
# report = TreasurerReport(startDate, endDate, newtable, initialBalance, currentBalance, deposits, depositsSum, withdrawals, withdrawalSum)
# filename = report.build()
# report.open()
